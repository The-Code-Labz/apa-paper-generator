"""
APA 7th Edition Paper Generator
FastAPI backend — generates a structured APA paper via Void AI (gpt-4o)
and exports a properly formatted .docx using python-docx.
"""
import os
import io
import json
import re
import logging
import datetime
import uuid

import httpx
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("apa")

# ------------------------------------------------------------------ config
VOID_API_BASE = os.environ.get("VOID_API_BASE", "https://api.voidai.app/v1")
VOID_API_KEY = os.environ.get("SHARED_VOID_API_FREE_KEY", "").strip()
MODEL = os.environ.get("APA_MODEL", "gpt-4o")
HERE = os.path.dirname(os.path.abspath(__file__))

app = FastAPI(title="APA Paper Generator")

# Directory for generated .docx files
os.makedirs("/tmp/apa-papers/", exist_ok=True)

# ------------------------------------------------------------------ models
class PaperRequest(BaseModel):
    title: str
    author: str
    course: str
    instructor: str
    institution: str = "Grand Canyon University"
    due_date: str
    prompt: str
    word_count: int = 1000
    sources: int = 3
    tone: str = "Academic"


class Section(BaseModel):
    heading: str
    body: str


class PaperResponse(BaseModel):
    title_page: dict
    abstract: str
    sections: list
    references: list


# ------------------------------------------------------------------ AI call
def build_system_prompt(req: PaperRequest) -> str:
    return f"""You are an expert academic writer producing papers in strict APA 7th edition (student paper) format.

Produce a COMPLETE academic paper based on the user's assignment prompt. Requirements:

- Target length: approximately {req.word_count} words across the body (introduction + body sections + conclusion). Aim to MEET this target — write substantive, fully developed paragraphs. It is better to slightly exceed {req.word_count} words than to fall short. Each body section should contain multiple developed paragraphs.
- Tone: {req.tone}.
- Include an Abstract of 150-250 words.
- Body MUST contain an Introduction, then 2 to 4 distinct body sections each with a clear APA Level 1 heading (title case, centered, bold — you only provide the heading text), then a Conclusion.
- Use proper APA 7th in-text citations in (Author, Year) format throughout the body where claims are made.
- Provide exactly {req.sources} references. They must be plausible, realistic academic sources (journal articles, books, or credible reports) formatted in full APA 7th edition reference style. Each in-text citation must correspond to one of the references.
- NO bullet points or lists in the body — full flowing paragraphs only.
- Do NOT include the words "Abstract", "Introduction", "Conclusion", or "References" as part of the body text fields; structure handles labels. The Introduction section heading should be the paper context (use the paper title concept) NOT the literal word "Introduction" unless natural — but typically the first section after the abstract is the Introduction with no heading in APA. To keep it simple: the FIRST section in the sections array is the Introduction (give it heading "Introduction"), the LAST is the Conclusion (heading "Conclusion").

Return ONLY valid JSON (no markdown fences) with this exact schema:
{{
  "title_page": {{
    "title": "{req.title}",
    "author": "{req.author}",
    "course": "{req.course}",
    "instructor": "{req.instructor}",
    "institution": "{req.institution}",
    "due_date": "{req.due_date}"
  }},
  "abstract": "string (150-250 words, single paragraph)",
  "sections": [
    {{ "heading": "Introduction", "body": "full paragraphs..." }},
    {{ "heading": "Some Level 1 Heading", "body": "full paragraphs..." }},
    {{ "heading": "Conclusion", "body": "full paragraphs..." }}
  ],
  "references": [
    "Full APA 7th edition reference string 1",
    "Full APA 7th edition reference string 2"
  ]
}}
"""


def build_user_prompt(req: PaperRequest) -> str:
    return f"""Assignment prompt / topic:
\"\"\"{req.prompt}\"\"\"

Paper title: {req.title}
Write the paper now and return the JSON object only."""


def extract_json(text: str) -> dict:
    text = text.strip()
    # strip code fences if any
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text).strip()
    # find first { ... last }
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON object found in model output")
    return json.loads(text[start:end + 1])


async def generate_paper(req: PaperRequest) -> dict:
    if not VOID_API_KEY:
        raise HTTPException(500, "Server missing AI API key (SHARED_VOID_API_FREE_KEY).")
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": build_system_prompt(req)},
            {"role": "user", "content": build_user_prompt(req)},
        ],
        "temperature": 0.7,
        "max_tokens": 6000,
        "response_format": {"type": "json_object"},
    }
    headers = {
        "Authorization": f"Bearer {VOID_API_KEY}",
        "Content-Type": "application/json",
    }
    async with httpx.AsyncClient(timeout=180) as client:
        # First attempt with response_format; fall back without it if rejected
        for attempt in range(2):
            try:
                r = await client.post(f"{VOID_API_BASE}/chat/completions",
                                      json=payload, headers=headers)
                if r.status_code >= 400:
                    log.warning("AI %s: %s", r.status_code, r.text[:300])
                    if attempt == 0:
                        payload.pop("response_format", None)
                        continue
                    raise HTTPException(502, f"AI provider error {r.status_code}")
                data = r.json()
                content = data["choices"][0]["message"]["content"]
                return extract_json(content)
            except (httpx.HTTPError,) as e:
                if attempt == 0:
                    continue
                raise HTTPException(502, f"AI request failed: {e}")
            except ValueError as e:
                if attempt == 0:
                    payload.pop("response_format", None)
                    continue
                raise HTTPException(502, f"Could not parse AI output: {e}")
    raise HTTPException(502, "AI generation failed")


# ------------------------------------------------------------------ docx
def style_normal(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def add_para(doc, text, *, align=None, bold=False, italic=False, indent=False, center=False):
    p = doc.add_paragraph()
    if center or align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def build_docx(paper: dict) -> io.BytesIO:
    doc = Document()
    # margins 1"
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style_normal(doc)

    tp = paper.get("title_page", {})

    # ---- Title page (student version) ----
    # ~3-4 blank double-spaced lines push title into upper-middle third
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, tp.get("title", ""), center=True, bold=True)
    doc.add_paragraph()  # blank line
    add_para(doc, tp.get("author", ""), center=True)
    add_para(doc, tp.get("institution", ""), center=True)
    add_para(doc, tp.get("course", ""), center=True)
    add_para(doc, tp.get("instructor", ""), center=True)
    add_para(doc, tp.get("due_date", ""), center=True)

    # ---- Abstract page ----
    doc.add_page_break()
    add_para(doc, "Abstract", center=True, bold=True)
    add_para(doc, paper.get("abstract", ""), indent=False)

    # ---- Body ----
    doc.add_page_break()
    # repeat title at top of body (APA student paper)
    add_para(doc, tp.get("title", ""), center=True, bold=True)
    for i, sec in enumerate(paper.get("sections", [])):
        heading = sec.get("heading", "").strip()
        body = sec.get("body", "").strip()
        # First section (Introduction) — APA convention: no heading, title serves
        if i == 0 and heading.lower() == "introduction":
            pass  # skip heading per APA
        else:
            add_para(doc, heading, center=True, bold=True)
        for para in re.split(r"\n\s*\n", body):
            para = para.strip()
            if para:
                add_para(doc, para, indent=True)

    # ---- References ----
    doc.add_page_break()
    add_para(doc, "References", center=True, bold=True)
    for ref in paper.get("references", []):
        ref = ref.strip()
        if not ref:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



# ------------------------------------------------------------------ markdown / agent helpers
def paper_to_markdown(paper: dict) -> str:
    tp = paper.get("title_page", {})
    lines = [
        f"# {tp.get('title', '')}",
        "",
        f"**{tp.get('author', '')}**",
        tp.get("institution", ""),
        tp.get("course", ""),
        tp.get("instructor", ""),
        tp.get("due_date", ""),
        "",
        "---",
        "",
        "## Abstract",
        paper.get("abstract", ""),
    ]
    for sec in paper.get("sections", []):
        lines.extend(["", "---", "", f"## {sec.get('heading', '')}", sec.get("body", "")])
    lines.extend(["", "---", "", "## References", ""])
    for ref in paper.get("references", []):
        lines.append(ref)
    return "\n".join(lines)

# ------------------------------------------------------------------ routes
@app.get("/api/health")
async def health():
    return {"ok": True, "model": MODEL, "key": bool(VOID_API_KEY)}


@app.post("/api/generate")
async def api_generate(req: PaperRequest):
    paper = await generate_paper(req)
    # basic validation / normalization
    paper.setdefault("title_page", {})
    paper["title_page"].setdefault("title", req.title)
    paper["title_page"].setdefault("author", req.author)
    paper["title_page"].setdefault("course", req.course)
    paper["title_page"].setdefault("instructor", req.instructor)
    paper["title_page"].setdefault("institution", req.institution)
    paper["title_page"].setdefault("due_date", req.due_date)
    paper.setdefault("abstract", "")
    paper.setdefault("sections", [])
    paper.setdefault("references", [])
    return JSONResponse(paper)


@app.post("/api/docx")
async def api_docx(paper: dict):
    buf = build_docx(paper)
    title = (paper.get("title_page", {}) or {}).get("title", "APA_Paper")
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", title)[:60] or "APA_Paper"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}.docx"'},
    )



@app.post("/api/agent/generate")
async def api_agent_generate(req: PaperRequest):
    paper = await generate_paper(req)
    # basic validation / normalization
    paper.setdefault("title_page", {})
    paper["title_page"].setdefault("title", req.title)
    paper["title_page"].setdefault("author", req.author)
    paper["title_page"].setdefault("course", req.course)
    paper["title_page"].setdefault("instructor", req.instructor)
    paper["title_page"].setdefault("institution", req.institution)
    paper["title_page"].setdefault("due_date", req.due_date)
    paper.setdefault("abstract", "")
    paper.setdefault("sections", [])
    paper.setdefault("references", [])

    paper_id = str(uuid.uuid4())
    buf = build_docx(paper)
    docx_path = os.path.join("/tmp/apa-papers/", f"{paper_id}.docx")
    with open(docx_path, "wb") as f:
        f.write(buf.getvalue())

    return JSONResponse({
        "paper": paper,
        "markdown": paper_to_markdown(paper),
        "download_url": f"/api/download/{paper_id}",
        "paper_id": paper_id,
    })


@app.get("/api/download/{paper_id}")
async def api_download(paper_id: str):
    base_dir = os.path.abspath("/tmp/apa-papers/")
    target = os.path.abspath(os.path.join(base_dir, f"{paper_id}.docx"))
    if not target.startswith(base_dir + os.sep) or not os.path.isfile(target):
        raise HTTPException(404, "File not found")
    return FileResponse(
        target,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{paper_id}.docx",
    )

@app.get("/")
async def index():
    return FileResponse(os.path.join(HERE, "static", "index.html"))


app.mount("/static", StaticFiles(directory=os.path.join(HERE, "static")), name="static")
